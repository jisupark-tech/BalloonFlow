"""
BalloonFlow — YAML 기획서 → .docx 문서 변환기 v5
현재 YAML 구조에 맞춰 전면 재작성.
6개 문서: 게임기획서, 시스템설계서, 밸런스설계서, 콘텐츠설계서, BM/LiveOps설계서, 빌드오더
"""
import io, sys, os, yaml
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"E:\AI\projects\BalloonFlow\design_workflow"
OUT = r"E:\AI\projects\BalloonFlow\docs"

# ── Helpers ──

def load_yaml(path):
    with open(os.path.join(BASE, path), 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def s(val):
    """Safe string conversion: strip trailing whitespace from YAML block scalars."""
    if val is None:
        return ""
    return str(val).strip()

def make_doc(title, subtitle=""):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    for level in [1, 2, 3]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Arial'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Arial'

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    return doc

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '4472C4'
        })
        shading.append(bg)
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(str(text), style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p

def add_info(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(s(value))
    run2.font.size = Pt(10)

def save(doc, filename):
    path = os.path.join(OUT, filename)
    doc.save(path)
    print(f"  Created: {filename}")


# ════════════════════════════════════════
# Doc 1: 게임 기획서 (Game Design Document)
# ════════════════════════════════════════
def gen_game_design():
    concept = load_yaml("concept.yaml")
    gd = load_yaml("layer1/game_design.yaml")

    doc = make_doc(
        "Balloon Flow!",
        "Dart × Balloon — Oddly Satisfying Queue Puzzle\nGame Design Document v2.0"
    )

    # 1. Vision
    doc.add_heading("1. 비전 (Vision)", level=1)
    add_info(doc, "한 줄 요약", concept['core_fun']['one_liner'])
    doc.add_paragraph(s(concept['core_fun']['extended']))

    doc.add_heading("Satisfaction Keywords", level=2)
    for kw in concept['core_fun'].get('satisfaction_keywords', []):
        add_bullet(doc, kw)

    # 2. Design Pillars
    doc.add_heading("2. Design Pillars (5개)", level=1)
    for pillar in concept['design_pillars']:
        doc.add_heading(f"Pillar {pillar['priority']}: {pillar['name']}", level=2)
        doc.add_paragraph(s(pillar['description']))

    # 3. Target Audience
    doc.add_heading("3. 타겟 유저", level=1)
    ta = concept['target_audience']
    add_table(doc,
        ["구분", "연령", "성별", "특성"],
        [
            ["Primary", ta['primary']['age_core'], ta['primary']['gender'], s(ta['primary']['description'])],
            ["Secondary", ta['secondary']['age'], "-", s(ta['secondary']['description'])],
        ])

    # 4. Core Mechanic
    doc.add_heading("4. 코어 메커닉", level=1)
    cm = concept['core_mechanic']
    add_info(doc, "타입", cm.get('basis', ''))
    add_info(doc, "타격 구조", cm.get('hit_structure', ''))
    add_info(doc, "레일 타입", cm.get('rail_type', ''))

    doc.add_heading("플레이 플로우", level=2)
    pf = cm.get('play_flow', {})
    for i in range(1, 7):
        step = pf.get(f'step_{i}', '')
        if step:
            add_bullet(doc, f"Step {i}: {s(step)}")

    doc.add_heading("Pixel Flow 대비 차별점", level=2)
    diff = cm.get('pixel_flow_difference', {})
    for key in ['feedback', 'structure', 'prototype_risk']:
        if key in diff:
            add_info(doc, key, diff[key])

    # 5. Core Loop
    doc.add_heading("5. 코어 루프", level=1)
    doc.add_heading("Inner Loop (한 판)", level=2)
    doc.add_paragraph(s(concept['core_loop']['inner']))
    doc.add_heading("Outer Loop (성장)", level=2)
    doc.add_paragraph(s(concept['core_loop']['outer']))

    doc.add_heading("실패 조건", level=2)
    fc = concept['core_loop']['fail_condition']
    add_info(doc, "상태", fc['status'])
    doc.add_paragraph(s(fc['primary']))

    # 5-1. Gimmicks
    doc.add_heading("5-1. 기믹 (5종)", level=1)
    gimmicks = gd.get('gimmicks', {})
    if gimmicks:
        doc.add_paragraph(s(gimmicks.get('description', '')))
        gimmick_list = gimmicks.get('list', [])
        rows = []
        for g in gimmick_list:
            levers = ", ".join(g.get('difficulty_levers', []))
            rows.append([g['name'], g.get('unlock', ''), s(g['description']), levers])
        add_table(doc, ["기믹", "해금", "설명", "난이도 레버"], rows)

    # 6. Systems Overview
    doc.add_heading("6. 시스템 구성", level=1)
    systems = gd['systems']
    rows = []
    for domain_name, info in systems.items():
        if isinstance(info, dict):
            sys_list = info.get('systems', [])
            deferred = info.get('deferred_systems', [])
            status = info.get('status', 'confirmed')
            if isinstance(sys_list, list):
                sys_str = ", ".join(sys_list)
                count = len(sys_list)
            else:
                sys_str = str(sys_list)
                count = 0
            rows.append([domain_name, status, s(info.get('description', '')), count, sys_str])
    add_table(doc, ["도메인", "상태", "설명", "수", "시스템 목록"], rows)

    # System Count
    sc = gd.get('system_count', {})
    confirmed = sc.get('confirmed', {})
    add_info(doc, "총 확정 시스템", confirmed.get('total_confirmed', ''))

    # 7. Art Direction
    doc.add_heading("7. 아트 디렉션", level=1)
    ad = concept['art_direction']
    add_info(doc, "스타일", ad['style'])
    add_info(doc, "방향", ad['direction'])
    add_info(doc, "풍선", ad['balloon_texture'])
    add_info(doc, "보관함", ad['holder_design'])
    add_info(doc, "다트", ad['dart_design'])
    add_info(doc, "배경", ad['background'])

    doc.add_heading("피드백 연출", level=2)
    fj = ad.get('feedback_juiciness', {})
    for key, val in fj.items():
        add_info(doc, key, val)

    # 8. Technical
    doc.add_heading("8. 기술 요구사항", level=1)
    tech = gd['technical']
    add_table(doc,
        ["항목", "값"],
        [
            ["Min Android SDK", tech['min_android_sdk']],
            ["Min iOS", tech['min_ios']],
            ["Target FPS", tech['target_fps']],
            ["Max Draw Calls", tech['max_draw_calls']],
            ["Max Balloon On Screen", tech['max_balloon_on_screen']],
            ["Object Pool", "필수" if tech['object_pool_required'] else "선택"],
            ["Orientation", tech['screen_orientation']],
            ["Render Pipeline", tech.get('render_pipeline', '')],
        ])

    # 9. References
    doc.add_heading("9. 레퍼런스 게임", level=1)
    refs = concept['references']
    add_table(doc,
        ["게임", "관련성", "차별점"],
        [[r['name'], r['relevance'], r.get('difference', '')] for r in refs])

    # 10. Risks
    doc.add_heading("10. 리스크", level=1)
    risks = concept.get('risks', [])
    add_table(doc,
        ["ID", "리스크", "대응"],
        [[r['id'], s(r['risk']), s(r['mitigation'])] for r in risks])

    save(doc, "BalloonFlow_게임기획서.docx")


# ════════════════════════════════════════
# Doc 2: 시스템 설계서 (System Spec)
# ════════════════════════════════════════
def gen_system_spec():
    spec = load_yaml("layer2/system_spec.yaml")

    doc = make_doc(
        "BalloonFlow 시스템 설계서",
        f"{spec['total_confirmed_systems']}개 확정 시스템 상세 명세\nSystem Specification v2.0"
    )

    domain_order = ['InGame', 'Content', 'Balance', 'UX', 'OutGame', 'BM']

    for domain_name in domain_order:
        domain_data = spec.get(domain_name, {})
        if not domain_data or not isinstance(domain_data, dict):
            continue

        doc.add_heading(f"Domain: {domain_name}", level=1)

        for sys_name, sys_info in domain_data.items():
            if not isinstance(sys_info, dict):
                continue

            doc.add_heading(sys_name, level=2)
            add_info(doc, "Layer", sys_info.get('layer', ''))
            add_info(doc, "Role", sys_info.get('role', ''))
            doc.add_paragraph(s(sys_info.get('description', '')))

            # Provides
            provides = sys_info.get('provides', [])
            if provides:
                doc.add_heading("Provides", level=3)
                for p in provides:
                    add_bullet(doc, p)

            # Requires
            requires = sys_info.get('requires', [])
            if requires:
                doc.add_heading("Requires", level=3)
                for r in requires:
                    add_bullet(doc, r)

            # Events
            events = sys_info.get('events_published', [])
            if events:
                doc.add_heading("Events Published", level=3)
                for e in events:
                    add_bullet(doc, e)

            # Config
            config = sys_info.get('config', {})
            if config and isinstance(config, dict):
                doc.add_heading("Config", level=3)
                for k, v in config.items():
                    add_info(doc, k, v)

            # Price table (BoosterManager)
            price_table = sys_info.get('price_table', {})
            if price_table:
                doc.add_heading("Price Table", level=3)
                rows = []
                for pid, pinfo in price_table.items():
                    if isinstance(pinfo, dict):
                        rows.append([pid, pinfo.get('coin_price', ''), pinfo.get('unit', ''), s(pinfo.get('note', ''))])
                add_table(doc, ["부스터", "코인 가격", "수량", "비고"], rows)

            # Note
            note = sys_info.get('note', '')
            if note:
                add_info(doc, "Note", note)

    save(doc, "BalloonFlow_시스템설계서.docx")


# ════════════════════════════════════════
# Doc 3: 밸런스 설계서
# ════════════════════════════════════════
def gen_balance():
    dc = load_yaml("balance/difficulty_curve.yaml")
    eco = load_yaml("balance/economy.yaml")

    doc = make_doc(
        "BalloonFlow 밸런스 설계서",
        "난이도 곡선 + 코인 경제\nBalance Design v2.0"
    )

    # 1. Difficulty Score
    doc.add_heading("1. 난이도 공식 (Difficulty Score)", level=1)
    ds = dc['difficulty_score']
    doc.add_paragraph(s(ds.get('description', '')))
    doc.add_paragraph(s(ds['formula']))

    doc.add_heading("1-1. 가중치", level=2)
    w = ds['weights']
    add_table(doc,
        ["변수", "가중치", "설명"],
        [
            ["W_color", w['W_color'], "색상 수 (2~8) — 홀더 슬롯 압박"],
            ["W_count", w['W_count'], "풍선 수 (6~65) — 보드 복잡도"],
            ["W_gimmick", w['W_gimmick'], "기믹 가중치 — 기믹별 독립 난이도 기여"],
            ["W_return", w['W_return'], "홀더 복귀 예상율 — 원형 루프 핵심 긴장감"],
            ["W_skew", w['W_skew'], "색상 편중도 — 특정 색 과다 시 홀더 관리 난이도↑"],
        ])

    doc.add_heading("1-2. 기믹 가중치 테이블", level=2)
    gwt = ds.get('gimmick_weight_table', {})
    rows = [[k, v] for k, v in gwt.items() if k != 'note']
    add_table(doc, ["기믹", "가중치"], rows)
    note = gwt.get('note', '')
    if note:
        doc.add_paragraph(s(note))

    doc.add_heading("1-3. 홀더 복귀 예상율", level=2)
    rr = ds.get('estimated_return_rate', {})
    doc.add_paragraph(s(rr.get('description', '')))
    doc.add_paragraph(s(rr.get('formula', '')))

    doc.add_heading("1-4. 색상 편중도", level=2)
    cs = ds.get('color_skew', {})
    doc.add_paragraph(s(cs.get('description', '')))
    doc.add_paragraph(s(cs.get('formula', '')))

    doc.add_heading("1-5. 난이도 예시값", level=2)
    examples = ds.get('examples', [])
    add_table(doc,
        ["Level", "Colors", "Balloons", "Gimmick", "Return", "Skew", "Score"],
        [[e['level'], e['num_colors'], e['balloon_count'],
          e['gimmick_weight'], e['return_rate'], e['color_skew'], e['score']]
         for e in examples])

    # 2. Sawtooth
    doc.add_heading("2. 톱니 패턴 (Sawtooth Modulation)", level=1)
    saw = dc['sawtooth']
    doc.add_paragraph(s(saw['description']))
    mods = saw['modifiers']
    add_table(doc,
        ["Position", "0", "1", "2", "3", "4"],
        [["Modifier"] + [str(m) for m in mods]])

    # 3. Magazine Count
    doc.add_heading("3. 탄창 수 (Magazine Count)", level=1)
    mc = dc['magazine_count']
    doc.add_paragraph(s(mc['description']))
    doc.add_paragraph(s(mc['formula']))

    doc.add_heading("3-1. Purpose Buffer", level=2)
    pb = mc.get('purpose_buffer', {})
    add_table(doc, ["목적", "Buffer 배율"],
        [[k, v] for k, v in pb.items()])

    doc.add_heading("3-2. 탄창 예시값", level=2)
    examples = mc.get('examples', [])
    if examples:
        add_table(doc,
            ["Level", "Balloons", "Colors", "Holders", "Optimal", "Purpose", "Buffer", "Total", "Per Holder"],
            [[e.get('level',''), e.get('balloons',''), e.get('colors',''), e.get('holders',''),
              e.get('optimal',''), e.get('purpose',''), e.get('buffer',''),
              e.get('total_magazine',''), e.get('per_holder','')]
             for e in examples])

    # 4. Star Thresholds
    doc.add_heading("4. 스타 점수 기준", level=1)
    st = dc['star_thresholds']
    doc.add_paragraph(s(st['formula']))
    doc.add_paragraph(s(st.get('note', '')))

    # 5. Clear Rate Model
    doc.add_heading("5. 클리어율 예측 모델", level=1)
    crm = dc['clear_rate_model']
    doc.add_paragraph(s(crm['formula']))
    doc.add_paragraph(s(crm['description']))

    # 6. Fail Condition
    doc.add_heading("6. 실패 조건", level=1)
    fc = dc['fail_condition']
    doc.add_paragraph(s(fc['primary']))

    # 7. Per-Package Targets
    doc.add_heading("7. 패키지별 클리어율 목표", level=1)
    ppt = dc.get('per_package_targets', {})
    pkgs = ppt.get('packages', {})
    rows = []
    for pkg_name, info in pkgs.items():
        rows.append([pkg_name,
            f"{info['cr_range'][0]}~{info['cr_range'][1]}",
            info['cr_avg'],
            f"{info['attempts_range'][0]}~{info['attempts_range'][1]}",
            info['attempts_avg']])
    add_table(doc, ["패키지", "CR 범위", "CR 평균", "시도 범위", "시도 평균"], rows)

    # 8. Economy
    doc.add_heading("8. 코인 경제", level=1)
    doc.add_heading("8-1. 재화 구조", level=2)
    curr = eco['currency_structure']
    add_info(doc, "주 재화", f"{curr['primary']['name']} ({curr['primary']['type']}) — 초기 {curr['primary']['initial']}개")
    add_info(doc, "보조 재화", f"{curr['secondary']['name']} — {curr['secondary']['status']}")

    doc.add_heading("8-2. 코인 소스 (획득)", level=2)
    cs = eco['coin_source']
    add_table(doc, ["소스", "코인"],
        [
            ["레벨 클리어 (노말)", cs['level_clear']['normal']],
            ["레벨 클리어 (하드)", cs['level_clear']['hard']],
            ["레벨 클리어 (슈퍼하드)", cs['level_clear']['super_hard']],
            ["보상형 광고", cs.get('rewarded_ad', 'TBD')],
            ["데일리 보상", cs.get('daily_reward', '')],
        ])

    doc.add_heading("8-3. 코인 싱크 (소비)", level=2)
    sink = eco['coin_sink']
    boosters = sink.get('boosters', {})
    rows = []
    for bid, binfo in boosters.items():
        rows.append([bid, binfo['cost'], binfo.get('unit', 1), s(binfo.get('note', ''))])
    add_table(doc, ["부스터", "가격", "수량", "비고"], rows)

    cont = sink.get('continue', {})
    add_info(doc, "컨티뉴", f"첫 1회 무료, 이후 {cont.get('escalating_cost', [])} (최대 {cont.get('max_uses_per_play', 3)}회)")
    add_info(doc, "하트 충전", f"{sink.get('heart_refill', {}).get('cost', '')}코인")

    doc.add_heading("8-4. 패키지별 코인 흐름", level=2)
    flow = eco.get('coin_flow_per_package', {})
    estimates = flow.get('estimates', {})
    rows = []
    for pkg_name, info in estimates.items():
        rows.append([pkg_name, s(info.get('source_estimate', '')),
                     s(info.get('sink_estimate', '')), s(info.get('net_flow', ''))])
    add_table(doc, ["패키지", "소스 예상", "싱크 예상", "순 흐름"], rows)

    doc.add_heading("8-5. 밸런스 원칙", level=2)
    for principle in eco.get('balance_principles', []):
        add_bullet(doc, principle)

    save(doc, "BalloonFlow_밸런스설계서.docx")


# ════════════════════════════════════════
# Doc 4: 콘텐츠 설계서
# ════════════════════════════════════════
def gen_content():
    bc = load_yaml("content/beat_chart.yaml")
    prog = load_yaml("content/progression.yaml")
    ld = load_yaml("content/level_design.yaml")

    doc = make_doc(
        "BalloonFlow 콘텐츠 설계서",
        "레벨 디자인 + Beat Chart + 진행 스케줄\nContent Design v2.0"
    )

    # 1. Beat Chart Overview
    doc.add_heading("1. Beat Chart 개요", level=1)
    gr = bc['global_rules']
    add_info(doc, "총 레벨", gr['total_levels'])
    add_info(doc, "패키지", f"{gr['packages']}개 × {gr['levels_per_package']}레벨")
    add_info(doc, "홀더 슬롯", gr['holder_slots'])
    add_info(doc, "실패 조건", gr['fail_condition'])

    doc.add_heading("난이도 원칙", level=2)
    dp = gr.get('difficulty_principles', {})
    for key in ['sawtooth_pattern', 'strict_ordering', 'new_element_intro_rule']:
        if key in dp:
            add_info(doc, key, dp[key])

    # 2. Per-Package Beat Chart
    doc.add_heading("2. 패키지별 Beat Chart", level=1)
    pkgs = bc['packages']
    for pkg_name, pkg in pkgs.items():
        doc.add_heading(f"{pkg_name}: {pkg['intent']}", level=2)
        doc.add_paragraph(s(pkg['description']))

        # Composition
        comp = pkg.get('composition', {})
        add_table(doc, ["목적", "수"],
            [[k, v] for k, v in comp.items()])

        # Color distribution
        cd = pkg.get('color_distribution', '')
        if cd:
            add_info(doc, "색상 분포", cd)
            add_info(doc, "색상 편중 의도", pkg.get('color_ratio_intent', ''))

        # Objects
        obj = pkg.get('objects', {})
        if obj:
            add_info(doc, "오브젝트 수", f"min {obj.get('min','')} / avg {obj.get('avg','')} / max {obj.get('max','')}")

        # CR targets
        cr = pkg.get('target_cr', {})
        att = pkg.get('target_attempts', {})
        if cr:
            add_info(doc, "클리어율", f"min {cr.get('min','')} / avg {cr.get('avg','')} / max {cr.get('max','')}")
        if att:
            add_info(doc, "시도 횟수", f"min {att.get('min','')} / avg {att.get('avg','')} / max {att.get('max','')}")

        # New elements
        ne = pkg.get('gimmicks', {}).get('new_elements', [])
        if ne:
            doc.add_heading("신규 요소", level=3)
            for elem in ne:
                add_bullet(doc, f"pos {elem['pos']}: {elem['element']} — {elem.get('note', '')}")

    # 3. Color Schedule
    doc.add_heading("3. 색상 도입 스케줄", level=1)
    colors = prog.get('color_schedule', [])
    add_table(doc,
        ["Level", "색상 수", "추가"],
        [[c['level'], c['colors'], c['added']] for c in colors])

    # 4. Gimmick Schedule
    doc.add_heading("4. 기믹 도입 스케줄", level=1)
    gimmicks = prog.get('gimmick_schedule', [])
    add_table(doc,
        ["Level", "기믹", "패키지"],
        [[g['level'], g['gimmick'], g['package']] for g in gimmicks])

    # 5. Level Purpose Assignment
    doc.add_heading("5. 레벨 목적 분류", level=1)
    pa = ld.get('purpose_assignment', {})
    doc.add_paragraph(s(pa.get('rule', '')))

    dt = pa.get('distribution_target', {})
    if dt:
        add_table(doc, ["목적", "비율"],
            [[k, v] for k, v in dt.items()])

    # 6. Tutorial Levels
    doc.add_heading("6. 튜토리얼 레벨 (1~5)", level=1)
    tut = ld.get('tutorial_levels', {})
    for lvl_name, info in tut.items():
        doc.add_heading(lvl_name.replace('_', ' ').title(), level=2)
        add_info(doc, "풍선", info.get('balloons', ''))
        add_info(doc, "색상", info.get('colors', ''))
        add_info(doc, "보관함", info.get('holders', ''))
        add_info(doc, "탄창", info.get('magazine_each', ''))
        add_info(doc, "가이드", info.get('guide', ''))
        add_info(doc, "목표", info.get('objective', ''))

    # 7. Gimmick Frequency
    doc.add_heading("7. 기믹 빈도 가이드", level=1)
    gfg = bc.get('gimmick_frequency_guidelines', {})
    plt = gfg.get('per_level_type', {})
    if plt:
        add_table(doc, ["레벨 타입", "최소 기믹", "최대 기믹"],
            [[k, v['min_gimmick_types'], v['max_gimmick_types']] for k, v in plt.items()])

    pgr = gfg.get('per_gimmick_ratio', {})
    if pgr:
        doc.add_heading("기믹별 오브젝트 비율 (%)", level=2)
        rows = []
        for gname, info in pgr.items():
            if gname == 'description':
                continue
            if isinstance(info, dict):
                rows.append([gname, info.get('min',''), info.get('avg',''), info.get('max','')])
        add_table(doc, ["기믹", "min%", "avg%", "max%"], rows)

    # 8. Day Progression
    doc.add_heading("8. 일자별 진행 예상", level=1)
    dp = prog.get('day_progression', {})
    est = dp.get('estimates', {})
    if est:
        rows = []
        for day in ['D1', 'D7', 'D14', 'D30']:
            casual = est.get('casual', {}).get(day, '')
            core = est.get('core', {}).get(day, '')
            rows.append([day, casual, core])
        add_table(doc, ["Day", "캐주얼 유저", "코어 유저"], rows)

    # 9. Longevity
    doc.add_heading("9. 콘텐츠 수명", level=1)
    lon = prog.get('longevity', {})
    add_info(doc, "런칭 콘텐츠", lon.get('launch_content', ''))
    add_info(doc, "업데이트 계획", lon.get('update_plan', ''))
    for plan in lon.get('post_100_plan', []):
        add_bullet(doc, plan)

    save(doc, "BalloonFlow_콘텐츠설계서.docx")


# ════════════════════════════════════════
# Doc 5: BM/LiveOps 설계서
# ════════════════════════════════════════
def gen_bm_liveops():
    bm = load_yaml("systems/bm.yaml")
    mon = load_yaml("bm/monetization.yaml")
    liveops = load_yaml("liveops/operations.yaml")
    concept = load_yaml("concept.yaml")

    doc = make_doc(
        "BalloonFlow BM & LiveOps 설계서",
        "수익화 구조 + 라이브 운영 (컨셉 확정분)\nBM & LiveOps v2.0"
    )

    # 1. BM Overview
    doc.add_heading("1. 수익화 모델", level=1)
    add_info(doc, "모델", mon.get('monetization_model', ''))
    add_info(doc, "범위", mon.get('scope', ''))
    add_info(doc, "BM 상태", bm.get('status', ''))

    # 2. Value System
    doc.add_heading("2. 가치 계산 체계", level=1)
    vs = mon.get('value_system', {})
    add_info(doc, "기준점", s(vs.get('base_rate', {}).get('reference', '')))
    add_info(doc, "코인/$", vs.get('base_rate', {}).get('coins_per_dollar', ''))

    doc.add_heading("설계 원칙", level=2)
    for p in vs.get('design_principles', []):
        add_bullet(doc, p)

    doc.add_heading("부스터 단위 가치", level=2)
    buv = vs.get('booster_unit_values', {})
    rows = []
    for bid, info in buv.items():
        if isinstance(info, dict):
            rows.append([bid, info.get('coins', ''), f"${info.get('usd', '')}"])
    add_table(doc, ["부스터", "코인 가치", "USD 가치"], rows)

    # 3. Coin Packs
    doc.add_heading("3. 표준 코인 상품 (6 Tiers)", level=1)
    cp = mon.get('coin_packs', {})
    table = cp.get('table', [])
    if table:
        add_table(doc,
            ["가격($)", "상품명", "코인", "효율", "이어하기 환산"],
            [[t['price'], t['name'], t['coins'], f"{t['total_efficiency']}x", t['continues']]
             for t in table])

    # 4. Bundles
    doc.add_heading("4. 표준 번들 상품 (5 Tiers)", level=1)
    bn = mon.get('bundles', {})
    table = bn.get('table', [])
    if table:
        add_table(doc,
            ["가격($)", "상품명", "코인", "부스터", "효율", "이어하기 환산"],
            [[t['price'], t['name'], t['coins'], t.get('boosters',''), f"{t['total_efficiency']}x", t['continues']]
             for t in table])

    # 5. Starter Pack
    doc.add_heading("5. 스타터팩", level=1)
    sp = mon.get('starter_pack', {})
    add_info(doc, "목적", sp.get('description', ''))
    add_info(doc, "해금", sp.get('unlock', ''))
    add_info(doc, "구매 제한", sp.get('purchase_limit', ''))
    table = sp.get('table', [])
    if table:
        add_table(doc,
            ["가격($)", "상품명", "코인", "부스터", "효율", "이어하기 환산"],
            [[t['price'], t['name'], t['coins'], t.get('boosters',''), f"{t['total_efficiency']}x", t['continues']]
             for t in table])

    # 6. Ad Removal
    doc.add_heading("6. 광고 제거", level=1)
    ar = mon.get('ad_removal', {})
    add_info(doc, "가격", ar.get('price', ''))
    add_info(doc, "효과", ar.get('effect', ''))

    # 7. Value Hierarchy
    doc.add_heading("7. 상품 간 가치 위계", level=1)
    vh = mon.get('value_hierarchy', [])
    if vh:
        add_table(doc, ["카테고리", "효율", "비고"],
            [[v['category'], v['efficiency'], v.get('note', '')] for v in vh])

    # 8. Ad Placements
    doc.add_heading("8. 광고 배치", level=1)
    ap = bm.get('ad_placements', {})
    if ap:
        doc.add_heading("보상형 광고", level=2)
        for r in ap.get('rewarded', []):
            add_bullet(doc, r)
        doc.add_heading("전면 광고", level=2)
        inter = ap.get('interstitial', {})
        add_info(doc, "보호 기간", inter.get('protection', ''))
        add_info(doc, "타이밍", inter.get('timing', ''))
        add_info(doc, "쿨다운", inter.get('cooldown', ''))

    # 9. Season Pass (DEFERRED)
    doc.add_heading("9. 시즌패스 (DEFERRED)", level=1)
    ssp = mon.get('season_pass', {})
    add_info(doc, "구조", ssp.get('structure', ''))
    add_info(doc, "무료 트랙", ssp.get('free_track', ''))
    add_info(doc, "프리미엄 트랙", ssp.get('premium_track', ''))
    add_info(doc, "철학", ssp.get('philosophy', ''))
    doc.add_paragraph("※ BM/LiveOps 디렉션 완료 후 확정 예정")

    # 10. LiveOps (DEFERRED)
    doc.add_heading("10. 라이브 운영 (DEFERRED — 컨셉 확정분만 기록)", level=1)

    doc.add_heading("10-1. 풍선 축제 (시즌 운영)", level=2)
    so = liveops.get('seasonal_operation', {})
    doc.add_paragraph(s(so.get('description', '')))
    for ex in so.get('examples', []):
        add_bullet(doc, ex)

    doc.add_heading("10-2. 연승 보너스", level=2)
    ws = liveops.get('win_streak_bonus', {})
    doc.add_paragraph(s(ws.get('description', '')))
    streaks = ws.get('streaks', [])
    if streaks:
        add_table(doc, ["연승", "보상"],
            [[st['streak'], st['reward']] for st in streaks])

    doc.add_heading("10-3. 데일리 챌린지 (오늘의 풍선)", level=2)
    dc = liveops.get('daily_challenge', {})
    doc.add_paragraph(s(dc.get('description', '')))
    for ex in dc.get('examples', []):
        add_bullet(doc, ex)
    add_info(doc, "보상", dc.get('reward', ''))

    doc.add_paragraph("※ 모든 LiveOps 시스템은 BM/LiveOps 디렉션 완료 후 구현 예정")

    save(doc, "BalloonFlow_BM_LiveOps설계서.docx")


# ════════════════════════════════════════
# Doc 6: 빌드 오더
# ════════════════════════════════════════
def gen_build_order():
    bo = load_yaml("layer2/build_order.yaml")

    doc = make_doc(
        "BalloonFlow 빌드 오더",
        f"{bo['confirmed_systems']}개 시스템 빌드 순서\nBuild Order v2.0"
    )

    # Phase 0 (unique structure: systems list)
    p0 = bo['phase_0']
    doc.add_heading(f"Phase 0: {p0['name']}", level=1)
    doc.add_paragraph(s(p0['description']))
    add_info(doc, "담당", p0['coder'])
    rows = []
    for sys in p0['systems']:
        rows.append([sys['name'], sys.get('layer', ''), sys.get('priority', ''), s(sys.get('description', ''))])
    add_table(doc, ["시스템", "Layer", "우선순위", "설명"], rows)

    # Phase 1-3 (parallel_assignment structure)
    for phase_key in ['phase_1', 'phase_2', 'phase_3']:
        phase = bo.get(phase_key)
        if not phase:
            continue
        doc.add_heading(f"{phase_key.replace('_', ' ').title()}: {phase['name']}", level=1)
        doc.add_paragraph(s(phase['description']))
        add_info(doc, "시스템 수", phase.get('systems_count', ''))
        add_info(doc, "의존성", phase.get('dependency', ''))

        pa = phase.get('parallel_assignment', {})
        for coder_name, systems in pa.items():
            doc.add_heading(coder_name, level=2)
            rows = []
            for sys in systems:
                rows.append([sys['name'], sys.get('layer', ''), s(sys.get('description', ''))])
            add_table(doc, ["시스템", "Layer", "설명"], rows)

    # Deferred
    doc.add_heading("Deferred Phase", level=1)
    deferred = bo.get('phase_deferred', {})
    add_info(doc, "트리거", deferred.get('trigger', ''))
    scope = deferred.get('expected_scope', {})
    for k, v in scope.items():
        add_info(doc, k, v)

    # Critical Path
    doc.add_heading("Critical Path", level=1)
    for path in bo.get('critical_path', []):
        add_bullet(doc, path)

    # Editor Scripts
    doc.add_heading("Editor Scripts", level=1)
    for es in bo.get('editor_scripts', []):
        add_bullet(doc, f"{es['name']} (Phase {es['phase']}): {es['description']}")

    save(doc, "BalloonFlow_빌드오더.docx")


# ════════════════════════════════════════
# MAIN
# ════════════════════════════════════════
if __name__ == "__main__":
    print("Generating BalloonFlow .docx documents (v5)...")
    gen_game_design()
    gen_system_spec()
    gen_balance()
    gen_content()
    gen_bm_liveops()
    gen_build_order()
    print(f"\nDone! All documents saved to: {OUT}")
